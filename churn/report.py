"""AI-written, board-ready churn report + export.

Privacy by design: only a compact, *aggregated* findings payload (schema +
statistics) is sent to the Anthropic API — never raw customer rows. This is
both a token-budget and a data-sensitivity decision (PE data).

Export: markdown always; HTML (with interactive figures) always; PDF via
weasyprint when its system libs are present; DOCX via python-docx when present.
"""

from __future__ import annotations

import os
import time
from dataclasses import dataclass, field
from typing import Any

import pandas as pd

DEFAULT_MODEL = "claude-opus-4-8"
MAX_TOKENS = 8000

# Models that support adaptive thinking + the effort parameter (used to deepen
# the analysis). Anything else falls back to a plain request.
_THINKING_MODELS = ("claude-opus-4-", "claude-sonnet-4-6", "claude-fable-5")

SYSTEM_PROMPT = """You are a senior data analyst writing a churn-analysis report \
for the business stakeholders who own the customer relationship (operations, \
customer success, product). Write clearly and concretely for a smart \
non-statistician. Lead with what the data shows and what to do about it.

The report must answer three questions, in this order, using markdown headings:

1. **Who churned** — which customer segments and profiles are leaving, and how \
   much more than the overall base rate. Use the segment churn rates and the \
   directional drivers to paint a concrete picture of the at-risk customer.
2. **Why they churned** — the ranked key drivers of churn, each in plain language \
   with its direction (raises/lowers churn) and the statistic that backs it \
   (odds ratio, SHAP importance, or statistical significance). Where SHAP is \
   provided, use it to explain *how* a driver moves churn.
3. **How to catch and prevent it** — split into two parts:
   a. *Identify earlier*: the leading indicators and early-warning signals an \
      analyst could monitor to flag an at-risk customer **before** they leave, \
      grounded in the drivers and the model's ability to rank risk (lift in the \
      top deciles → a watchlist / risk-scoring approach).
   b. *Prevent*: concrete interventions, each tied to a specific driver and \
      framed as a hypothesis to test, not a guaranteed fix.

Open with a short **Executive summary** (the base churn rate plus the 2-3 \
headline takeaways) and close with **Caveats**.

Structure:
- **Executive summary**
- **1. Who churned**
- **2. Why they churned**
- **3. Catching churn earlier & preventing it**
- **Caveats** — correlation is not causation; data limitations; class imbalance; \
  sample size.

HARD RULES:
- Use ONLY numbers that appear in the supplied findings payload. Do NOT invent \
  statistics, benchmarks, dollar figures, or industry comparisons.
- If you state an assumption, flag it explicitly as an assumption.
- A "driver" means associated + predictive, NOT causal. Never imply an \
  intervention is guaranteed to work.
- Do NOT frame this around private equity, valuation multiples, NRR/GRR, or \
  investor economics. The audience is the operating team.
"""


class ReportError(RuntimeError):
    pass


# ----------------------------------------------------------------------
# Findings payload (aggregates only — no raw rows)
# ----------------------------------------------------------------------
def build_findings_payload(
    *,
    dataset_name: str,
    n_rows: int,
    n_features: int,
    base_rate: float,
    cleaning_log: list[str],
    driver_parent_table: pd.DataFrame,
    stats_table: pd.DataFrame | None,
    model_metrics: dict[str, dict[str, Any]],
    segment_rates: dict[str, list[dict[str, Any]]] | None = None,
    top_n: int = 8,
) -> dict[str, Any]:
    """Assemble the compact structured payload the model writes from."""
    drivers = []
    for _, r in driver_parent_table.head(top_n).iterrows():
        d = {
            "feature": r.get("parent", r.get("feature")),
            "direction": r.get("direction"),
            "strength_rank": int(r.get("strength_rank", 0)),
            "significant": bool(r.get("any_significant", r.get("significant", False))),
        }
        if "shap_importance" in r and pd.notna(r["shap_importance"]):
            d["shap_importance"] = round(float(r["shap_importance"]), 4)
        drivers.append(d)

    stats_summary = []
    if stats_table is not None and not stats_table.empty:
        for _, r in stats_table.head(top_n).iterrows():
            stats_summary.append({
                "feature": r["feature"],
                "test": r["test"],
                "p_fdr": round(float(r["p_fdr"]), 5),
                "significant_fdr": bool(r["significant_fdr"]),
                "effect_size": round(float(r["effect_size"]), 3),
                "effect_name": r["effect_name"],
                "direction": r["direction"],
            })

    return {
        "dataset": {"name": dataset_name, "rows": n_rows, "features": n_features},
        "base_churn_rate": round(float(base_rate), 4),
        "cleaning_log": cleaning_log,
        "top_drivers": drivers,
        "univariate_stats": stats_summary,
        "model_metrics": model_metrics,
        "segment_churn_rates": segment_rates or {},
        "notes": [
            "Drivers are associations, not proven causes.",
            "No external benchmark was supplied; do not compare to industry figures.",
            "Risk-ranking deciles show how well the model concentrates churners into "
            "the highest-risk customers — the basis for early identification.",
        ],
    }


def metrics_for_payload(model_result) -> dict[str, Any]:
    """Compact metric dict for one model (CV mean±std + test + risk ranking)."""
    gains = model_result.gains
    risk_deciles = []
    if not gains.empty:
        for _, r in gains.head(3).iterrows():
            risk_deciles.append({
                "decile": int(r["decile"]),
                "churn_rate": round(float(r["churn_rate"]), 4),
                "lift": round(float(r["lift"]), 3),
                "cum_pct_churners_captured": round(float(r["cum_pct_churners_captured"]), 4),
            })
    return {
        "name": model_result.name,
        "cv": {k: {"mean": round(v[0], 4), "std": round(v[1], 4)} for k, v in model_result.cv.items()},
        "test": {k: round(v, 4) for k, v in model_result.test.items() if v == v},
        "top_decile_lift": round(float(gains["lift"].iloc[0]), 3) if not gains.empty else None,
        "risk_ranking_top_deciles": risk_deciles,
    }


# ----------------------------------------------------------------------
# Anthropic call
# ----------------------------------------------------------------------
@dataclass
class ReportConfig:
    model: str = field(default_factory=lambda: os.getenv("REPORT_MODEL", DEFAULT_MODEL))
    api_key: str | None = field(default_factory=lambda: os.getenv("ANTHROPIC_API_KEY"))
    max_tokens: int = MAX_TOKENS
    max_retries: int = 3


def generate_report(payload: dict[str, Any], config: ReportConfig | None = None) -> str:
    """Call Anthropic to write the memo. Raises :class:`ReportError` on failure."""
    config = config or ReportConfig()
    if not config.api_key:
        raise ReportError(
            "ANTHROPIC_API_KEY is not set. Add it to your .env (see .env.example) "
            "to generate the AI report. All other stages work without it."
        )
    try:
        import anthropic
    except ImportError as exc:  # pragma: no cover
        raise ReportError("The 'anthropic' package is not installed.") from exc

    import json

    client = anthropic.Anthropic(api_key=config.api_key, max_retries=config.max_retries)
    user_msg = (
        "Write the churn report from this findings payload (JSON). Use only these "
        "numbers.\n\n```json\n" + json.dumps(payload, indent=2, default=str) + "\n```"
    )

    # Deepen the analysis with adaptive thinking + high effort on capable models.
    extra: dict[str, Any] = {}
    if config.model.startswith(_THINKING_MODELS):
        extra = {"thinking": {"type": "adaptive"}, "output_config": {"effort": "high"}}

    last_err: Exception | None = None
    for attempt in range(config.max_retries):
        try:
            resp = client.messages.create(
                model=config.model,
                max_tokens=config.max_tokens,
                system=SYSTEM_PROMPT,
                messages=[{"role": "user", "content": user_msg}],
                **extra,
            )
            return "".join(b.text for b in resp.content if getattr(b, "type", "") == "text")
        except Exception as exc:  # noqa: BLE001 - map SDK errors to a clean message
            last_err = exc
            name = type(exc).__name__
            if "Authentication" in name or "PermissionDenied" in name:
                raise ReportError(f"Anthropic auth failed ({name}). Check ANTHROPIC_API_KEY.") from exc
            if attempt < config.max_retries - 1:
                time.sleep(1.5 * (attempt + 1))
                continue
    raise ReportError(f"Anthropic API call failed after retries: {last_err}")


# ----------------------------------------------------------------------
# Export
# ----------------------------------------------------------------------
def report_figure_subset(figures: dict[str, Any]) -> dict[str, Any]:
    """The driver-explaining figures to embed in the report, in narrative order.

    SHAP beeswarm + dependence plots explain *why* customers churn; the driver
    importance bar ranks the drivers; segment/distribution plots show *who*.
    Model-performance curves (ROC/PR/calibration) are left for the Visualize tab.
    """
    def _title(fig: Any, fallback: str) -> str:
        try:
            t = fig.layout.title.text
            return t if t else fallback
        except Exception:  # noqa: BLE001
            return fallback

    keys: list[str] = []
    for key in ("driver_importance", "shap_beeswarm"):
        if key in figures:
            keys.append(key)
    keys += [k for k in figures if k.startswith("shap_dependence::")]
    keys += [k for k in figures if k.startswith(("segment::", "distribution::"))]

    ordered: dict[str, Any] = {}
    for key in keys:
        fig = figures[key]
        ordered[_title(fig, key)] = fig
    return ordered


def export_markdown(md: str) -> bytes:
    return md.encode("utf-8")


def _md_to_html(md: str) -> str:
    import markdown as md_lib

    body = md_lib.markdown(md, extensions=["tables", "fenced_code"])
    return (
        "<!doctype html><html><head><meta charset='utf-8'>"
        "<style>body{font-family:-apple-system,Segoe UI,Roboto,sans-serif;"
        "max-width:820px;margin:40px auto;line-height:1.5;color:#1a1a1a;padding:0 20px}"
        "h1,h2{border-bottom:1px solid #eee;padding-bottom:4px}"
        "table{border-collapse:collapse}td,th{border:1px solid #ccc;padding:4px 8px}"
        "code{background:#f5f5f5;padding:1px 4px}</style></head><body>"
        + body + "</body></html>"
    )


def export_html(md: str, figures: dict[str, Any] | None = None) -> bytes:
    """HTML export with interactive Plotly figures appended."""
    html = _md_to_html(md)
    if figures:
        chunks = ["<hr><h1>Figures</h1>"]
        for i, (name, fig) in enumerate(figures.items()):
            chunks.append(f"<h2>{name}</h2>")
            chunks.append(fig.to_html(full_html=False, include_plotlyjs="cdn" if i == 0 else False))
        html = html.replace("</body></html>", "".join(chunks) + "</body></html>")
    return html.encode("utf-8")


def weasyprint_available() -> bool:
    try:
        import weasyprint  # noqa: F401

        return True
    except Exception:  # noqa: BLE001 - missing system libs
        return False


def export_pdf(md: str, figures: dict[str, Any] | None = None) -> bytes | None:
    """PDF via weasyprint. Returns None if weasyprint (system libs) is absent."""
    if not weasyprint_available():
        return None
    import weasyprint

    html = _md_to_html(md)
    image_html = _figures_as_png_html(figures) if figures else ""
    if image_html:
        html = html.replace("</body></html>", image_html + "</body></html>")
    return weasyprint.HTML(string=html).write_pdf()


def _figures_as_png_html(figures: dict[str, Any]) -> str:
    """Embed figures as static PNGs (needs kaleido). Skips silently if absent."""
    import base64

    try:
        chunks = ["<hr><h1>Figures</h1>"]
        for name, fig in figures.items():
            png = fig.to_image(format="png", width=820, height=480, scale=2)
            b64 = base64.b64encode(png).decode()
            chunks.append(f"<h2>{name}</h2><img style='max-width:100%' src='data:image/png;base64,{b64}'/>")
        return "".join(chunks)
    except Exception:  # noqa: BLE001 - kaleido not installed; PDF stays text-only
        return ""


def export_docx(md: str) -> bytes | None:
    """Lightweight DOCX export (headings + paragraphs). None if python-docx absent."""
    try:
        import io

        from docx import Document
    except Exception:  # noqa: BLE001
        return None

    doc = Document()
    for line in md.splitlines():
        stripped = line.strip()
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith(("- ", "* ")):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        elif stripped:
            doc.add_paragraph(stripped)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
